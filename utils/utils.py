import json
from pathlib import Path
import tempfile
import os

def extract_json_from_response(response: str) -> dict:
    """
    Extracts JSON from a response string that is wrapped in <json> or ```json tags.
    """

    response = response.replace("<json>", "").replace("</json>", "")
    response = response.replace("```json", "").replace("```", "")

    return json.loads(response)

def recuperar_contexto(clean_text:str, top_k:int, filtro:str) -> str:
    """Consulta pinecone para recuperar ejemplos de buenos requisitos
    y plantillas IEEE 830 relevantes al texto del usuario"""

    try:
        from rag.retriever import recuperar_contexto
        docs = recuperar_contexto(
            query = clean_text,
            top_k = top_k,
            filtro = {"agente": filtro}
        )

        return "\n\n".join(d.page_content for d in docs)
    except Exception as exc:
        return "" 

def _extraer_txt(ruta: str) -> str:
    """
    Extrae y devuelve el contenido de un archivo .txt.

    Args:
        ruta_archivo: Ruta del archivo TXT.
        encoding: Codificación del archivo (default utf-8).

    Returns:
        str: Texto contenido en el archivo.

    Raises:
        FileNotFoundError: Si el archivo no existe.
        ValueError: Si el archivo no es .txt.
    """

    ruta = Path(ruta)

    if not ruta.exists():
        raise FileNotFoundError(f"El archivo no existe: {ruta}")

    if ruta.suffix.lower() != ".txt":
        raise ValueError("El archivo debe tener extensión .txt")

    with open(ruta, "r", encoding="utf-8") as archivo:
        texto = archivo.read()

    return texto

# ── PDF ───────────────────────────────────────────────────────────────────────

def _extraer_pdf(ruta: str) -> str:
    """
    Extrae texto de un PDF usando pdfplumber.
    Preserva el orden de lectura página a página.
    pip install pdfplumber
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Instala pdfplumber: pip install pdfplumber")

    paginas: list[str] = []

    with pdfplumber.open(ruta) as pdf:
        for i, pagina in enumerate(pdf.pages, start=1):
            texto = pagina.extract_text()
            if texto:
                paginas.append(f"--- Página {i} ---\n{texto.strip()}")

    return "\n\n".join(paginas)


# ── Word ──────────────────────────────────────────────────────────────────────

def _extraer_word(ruta: str) -> str:
    """
    Extrae texto de un archivo .docx usando python-docx.
    Preserva párrafos y texto de tablas.
    pip install python-docx
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Instala python-docx: pip install python-docx")

    doc        = Document(ruta)
    fragmentos: list[str] = []

    # Párrafos
    for parrafo in doc.paragraphs:
        if parrafo.text.strip():
            fragmentos.append(parrafo.text.strip())

    # Texto en tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            celda_texts = [c.text.strip() for c in fila.cells if c.text.strip()]
            if celda_texts:
                fragmentos.append(" | ".join(celda_texts))

    return "\n\n".join(fragmentos)


# ── Email ─────────────────────────────────────────────────────────────────────

def _extraer_email(ruta: str) -> str:
    """
    Extrae texto de archivos .eml usando mail-parser.
    Incluye asunto, remitente y cuerpo del mensaje.
    pip install mail-parser
    """
    try:
        import mailparser
    except ImportError:
        raise ImportError("Instala mail-parser: pip install mail-parser")

    mail = mailparser.parse_from_file(ruta)

    partes: list[str] = []

    if mail.subject:
        partes.append(f"Asunto: {mail.subject}")
    if mail.from_:
        partes.append(f"De: {mail.from_}")
    if mail.to:
        partes.append(f"Para: {mail.to}")
    if mail.body:
        partes.append(f"Cuerpo:\n{mail.body.strip()}")

    return "\n\n".join(partes)


# ── Audio ─────────────────────────────────────────────────────────────────────

def _extraer_audio(ruta: str) -> str:
    """
    Transcribe audio usando la API de OpenAI Whisper.
    Soporta mp3, wav, m4a, ogg, webm.
    pip install openai
    Requiere: OPENAI_API_KEY en .env
    """
    try:
        from openai import AsyncOpenAI
    except ImportError:
        raise ImportError("Instala openai: pip install openai")

    client = AsyncOpenAI(api_key=os.environ["OPENAI_API_KEY"])

    with open(ruta, "rb") as audio_file:
        transcripcion = client.audio.transcriptions.create(
            model = "whisper-1",
            file  = audio_file,
            language = "es",        # forzar español; cambiar según idioma esperado
        )

    return transcripcion.text


# ── Video ─────────────────────────────────────────────────────────────────────

# def _extraer_video(ruta: str) -> str:
#     """
#     Extrae el audio de un video con moviepy y luego lo transcribe con Whisper.
#     Soporta mp4, mov, avi, mkv.
#     pip install moviepy openai
#     """
#     try:
#         from moviepy.editor import VideoFileClip
#     except ImportError:
#         raise ImportError("Instala moviepy: pip install moviepy")
    
#     # Extraer audio del video a un archivo temporal WAV
#     with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tmp:
#         ruta_audio_tmp = tmp.name

#     try:
#         clip = VideoFileClip(ruta)
#         clip.audio.write_audiofile(ruta_audio_tmp, verbose=False, logger=None)
#         clip.close()

#         # Reusar el extractor de audio
#         texto = _extraer_audio(ruta_audio_tmp)
#     finally:
#         Path(ruta_audio_tmp).unlink(missing_ok=True)

#     return texto


# ── Texto plano ────────────────────────────────────────────────────────────────

def _extraer_texto_plano(ruta: str) -> str:
    """Fallback para archivos de texto sin formato conocido."""
    with open(ruta, encoding="utf-8", errors="ignore") as f:
        return f.read()
    
def detectar_tipo(archivos: list[dict]) -> str:
    """
    Determina el tipo_entrada del AgentState según los archivos recibidos.
    Si hay mezcla de tipos, retorna 'documento' como valor genérico.
    """
    tipos = {a.get("tipo", "").lower() for a in archivos}

    if tipos <= {"mp3", "wav", "m4a", "audio"}:
        return "conversacion"
    if tipos <= {"mp4", "mov", "avi", "video"}:
        return "conversacion"
    return "documento"

def _extraer_texto(ruta: str, tipo: str) -> str:
    """
    Enruta la extracción de texto al extractor correcto según el tipo.
    """
    extractores = {
        "pdf":   _extraer_pdf,
        "word":  _extraer_word,
        "docx":  _extraer_word,
        "email": _extraer_email,
        "eml":   _extraer_email,
        "audio": _extraer_audio,
        "mp3":   _extraer_audio,
        "wav":   _extraer_audio,
        "m4a":   _extraer_audio,
        # "video": _extraer_video,
        # "mp4":   _extraer_video,
        # "mov":   _extraer_video,
        # "avi":   _extraer_video,
        "txt":   _extraer_txt,
    }

    extractor = extractores.get(tipo)
    if not extractor:
        print("Tipo '{}' no soportado — intentando como texto plano".format(tipo))
        return _extraer_texto_plano(ruta)

    return extractor(ruta)

def procesar_documentos(archivos:list[dict]) -> str:
    
    fragmentos: list[str] = []

    for archivo in archivos:
        ruta = archivo.get("ruta", "")
        tipo  = archivo.get("tipo", "").lower()
        nombre= archivo.get("nombre", ruta)

        try:
            texto = _extraer_texto(ruta, tipo)
            if texto:
                fragmentos.append(f"[{nombre}]\n{texto}")
                
                print("'{}' procesado".format(nombre))
            else:
                print("'{}' no produjo texto extraíble".format(nombre))

        except Exception as exc:
            print("Error procesando '{}': {}".format(nombre, exc))

    return "\n\n".join(fragmentos)

def indexar(archivos:list[dict], thread_id:str, project_name:str):
    
    from rag.indexer import indexar_documento

    fragmentos: list[str] = []

    for archivo in archivos:
        ruta = archivo.get("ruta", "")
        tipo  = archivo.get("tipo", "").lower()
        nombre= archivo.get("nombre", ruta)

        try:
            texto = _extraer_texto(ruta, tipo)
            if texto:
                fragmentos.append(f"[{nombre}]\n{texto}")
                
                indexar_documento(
                    texto     = texto,
                    metadata  = {
                        "thread_id":    thread_id,
                        "project_name": project_name,
                        "tipo":         tipo,
                    }
                )
                
                print("""'{}' indexado en Pinecone ({} chars) - 
                      nombre del proyecto: {}, 
                      thread_id: {}""".format(nombre, len(texto), project_name, thread_id))
            else:
                print("'{}' no produjo texto extraíble".format(nombre))

        except Exception as exc:
            print("Error procesando '{}': {}".format(nombre, exc))

def indexar_requisitos(requisitos_str, thread_id:str, project_name:str):
    from rag.indexer import indexar_documento

    try:
        if requisitos_str:
            indexar_documento(
                texto     = requisitos_str,
                metadata  = {
                    "thread_id":    thread_id,
                    "project_name": project_name
                }
            )
                
            print("""Requerimientos indexados en Pinecone - 
                      nombre del proyecto: {}, 
                      thread_id: {}""".format(project_name, thread_id))
        else:
            print("No hay requisitos_str")

    except Exception as exc:
        print("Error procesando {}".format(exc))